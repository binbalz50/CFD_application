from docx import Document
from docx.shared import Cm, Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import pandas as pd
import matplotlib.pyplot as plt

class form:
    def __init__(self):
        self.doc = Document()

    def title(self, title): #Viết tiêu đề
        text = self.doc.add_paragraph(title)
        run = text.runs[0]
        run.font.name = 'Times New Roman'
        run.font.size = Pt(16)
        text.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    def center_paragraph(self,text): #Viết đoạn văn căn giữa
        para = self.doc.add_paragraph(text)
        run = para.runs[0]
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Căn giữa

    def left_paragraph(self,text): #Viết đoạn văn căn giữa
        paragraph = self.doc.add_paragraph(text)
        run = paragraph.runs[0]
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Căn trái


    def add_image(self, image,sub): #Thêm ảnh căn giữa
        paragraph = self.doc.add_paragraph()
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = paragraph.add_run()
        run.add_picture(image, width=Inches(6))
        sub_caption = self.doc.add_paragraph(sub)
        sub_caption.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    def add_table(self,mach,aoa,temp,pressure):
        self.table=self.doc.add_table(rows=2,cols=5)
        self.table.style="Table Grid"
        headers = ["Thông số", "Số Mach", "Góc tấn(độ)", "Nhiệt độ(K)", "Áp suất(Pa)"]
        for col_idx, header in enumerate(headers):
            self.table.cell(0, col_idx).text = header  # Ghi tiêu đề vào hàng đầu tiên

        # Thêm dữ liệu mẫu vào hàng thứ hai
        data = ["Giá trị", mach, aoa, temp, pressure]
        for col_idx, value in enumerate(data):
            self.table.cell(1, col_idx).text = value  # Ghi dữ liệu vào hàng thứ hai
        
    def save(self):
        self.doc.save("bao_cao.docx")

    def plot(self,csv,image):
        self.read=pd.read_csv(csv, delimiter=",")
        self.read.columns = self.read.columns.str.replace('"', '').str.strip()
        plt.figure(figsize=(10, 5))
        plt.plot(self.read["Inner_Iter"], self.read["CL"], label="CL", color="blue", marker='o')
        plt.plot(self.read["Inner_Iter"], self.read["CD"], label="CD", color="red", marker='x') 
        plt.xlabel("Iteration")
        plt.ylabel("Coefficient Value")
        plt.title("CL & CD vs Iteration")
        plt.legend()
        plt.grid(True)
        plt.savefig(image, dpi=300, bbox_inches="tight") 

class export:
     def __init__(self, mach, aoa, temp, pressure,data,csv_path,mesh_path,field_path,plot_path,field):
        self.exp = form()
        self.exp.title(title=f"Kết quả mô phỏng khí động mô hình {data}")
        self.exp.left_paragraph(text="1.Điều kiện ban đầu")
        self.exp.add_table(mach, aoa, temp, pressure) #Thêm dữ liệu vào bảng
        self.exp.left_paragraph(text="2.Lưới tính toán")
        self.exp.add_image(image=mesh_path,sub=f"Lưới tính toán của mô hình {data}")
        self.exp.left_paragraph(text="3.Kết quả")
        self.exp.add_image(image=field_path,sub=f"Kết quả mô phỏng mô hình {data}")
        self.exp.plot(csv=csv_path,image=plot_path)
        self.exp.add_image(image=plot_path,sub=f"Đồ thị mô hình {data}")
        self.exp.save()

